from os import path
import unittest
import unicodedata
import re
from hypothesis import given, note, settings, strategies as st
from docx import Document
from extractors.docx_text_extractor import DocxTextExtractor
from extractors.html_text_extractor import HtmlTextExtractor
from readers.html_reader import RemoteHtmlReader
from extraction_objects.docx_extraction_object import DocxExtractionObject


class TestDocxTextExtractor(unittest.TestCase):
    def setUp(self):
        self.extractor = DocxTextExtractor()
        tests_dir_path = path.abspath(path.dirname(__file__))
        self.samples_dir = path.join(tests_dir_path, 'test_examples')

    def _read_file(self, path: str) -> DocxExtractionObject:
        return DocxExtractionObject(Document(path))

    def test_table_removal(self):
        filepath = path.join(self.samples_dir, 'table_text.docx')
        eo = self._read_file(filepath)
        self.assertEqual(
            self.extractor.extract(eo),
            'А это тестовый текст.'
        )

    def test_running_lines_removal(self):
        filepath = path.join(self.samples_dir, 'colontitle_text.docx')
        eo = self._read_file(filepath)
        self.assertEqual(
            self.extractor.extract(eo),
            'Это текст в тестовом документе с колонтитулами.'
        )

    def test_image_removal(self):
        filepath = path.join(self.samples_dir, 'picture_text.docx')
        eo = self._read_file(filepath)
        self.assertEqual(
            self.extractor.extract(eo),
            'Это текст в документе с пдописанным изображением'
        )

    def test_multiple_spaces_replacing(self):
        filepath = path.join(self.samples_dir, 'multi_spaces.docx')
        eo = self._read_file(filepath)
        self.assertEqual(
            self.extractor.extract(eo),
            'Это первое предлжение текста с несколькими пробелами в некоторых местах '
            'например в этом и ещё вот в этом.'
        )

    def test_multiple_breaklines_replacing(self):
        filepath = path.join(
            self.samples_dir, 'multiline_multi_line_breaks.docx'
        )
        eo = self._read_file(filepath)
        self.assertEqual(
            self.extractor.extract(eo),
            'And more text. And more text. And more text. And more text. And more\n'
            'text. And more text. Boring, zzzzz. And more text. And more text. And\n'
            'more text. And more text. And more text. And more text. And more text.'
        )

    def test_formula_removal(self):
        filepath = path.join(self.samples_dir, 'formula_text.docx')
        eo = self._read_file(filepath)
        self.assertEqual(
            self.extractor.extract(eo),
            'Это текст в файле с формулой'
        )

    def test_strings_trimming(self):
        filepath = path.join(self.samples_dir, 'tabulation.docx')
        eo = self._read_file(filepath)
        self.assertEqual(
            self.extractor.extract(eo),
            'And more text. And more text. And more text. And more text. '
            'And more text. And more text. Boring, zzzzz. And more text. '
            'And more text. And more text. And more text. And more text. '
            'And more text. And more text. And more text. And more text. '
            'And more text. And more text. And more text. And more text. '
            'Boring, zzzzz. And more text. And more text. And more text. '
            'And more text. And more text. And more text. And more text'
        )

    def test_empty_file(self):
        document = DocxExtractionObject(Document()) 
        self.assertEqual(self.extractor.extract(document), "")


    @settings(max_examples=50)
    @given(st.text(st.characters(blacklist_categories=("Cc", "Cs")), min_size=1))
    def test_random_characters_doc(self, s):
        doc = Document()
        doc.add_paragraph(s)
        doc_obj = DocxExtractionObject(doc)
        res = self.extractor.extract(doc_obj)
        note(f"Raw data: {s}")
        note(f"Result: {res}")
        norm_s = unicodedata.normalize("NFKC", s)
        #print(norm_s, res)
        self.assertEqual(res, self.extractor.clean_text(norm_s))


    def test_long_multilang(self):
        html_extr = HtmlTextExtractor()
        remote_reader = RemoteHtmlReader(headers={})
        html_obj = remote_reader.read("https://ru.wikipedia.org/wiki/%D0%92%D0%BE%D0%B9%D0%BD%D0%B0_%D0%B8_%D0%BC%D0%B8%D1%80#cite_note-31")
        test_text = html_extr.extract(html_obj)
        doc = Document()
        doc.add_paragraph(test_text)
        doc_obj = DocxExtractionObject(doc)
        res = self.extractor.extract(doc_obj)
        res_predicate = all([re.search(phrase, res) for phrase in [
            "Главная тема — историческая судьба русского народа в Отечественной войне 1812 года", 
            "La Guerre et la Paix", 
            "The 10 Greatest Books of All Time"
            ]
        ])
        self.assertTrue(res_predicate)

    def test_no_alphabetic_chars(self):
        doc = Document()
        test_text = '!"#$%&\'()*+,-./0123456789:;<=>?@[\]^_`{|}‘’‚‛“”„'
        doc.add_paragraph(test_text)
        obj = DocxExtractionObject(doc)
        self.assertEqual(
            self.extractor.extract(obj),
            test_text
        )

    def test_latin_alphabet(self):
        doc = Document()
        capital_alphabet = 'ABCDEFGHIJKLMNOPQRSTUVWXYZ'
        alphabet = 'abcdefghijklmnopqrstuvwxyz'
        text = f'{capital_alphabet}\n{alphabet}'
        doc.add_paragraph(text)
        obj = DocxExtractionObject(doc)
        self.assertEqual(
            self.extractor.extract(obj),
            text
        )

    def test_cyrillic_alphabet(self):
        doc = Document()
        capital_alphabet = 'АБВГДЕЁЖЗИЙКЛМНОПРСТУФХЦЧШЩЪЫЬЭЮЯ'
        alphabet = 'абвгдеёжзийклмнопрстуфхцчшщъыьэюя'
        text = f'{capital_alphabet}\n{alphabet}'
        doc.add_paragraph(text)
        obj = DocxExtractionObject(doc)
        self.assertEqual(
            self.extractor.extract(obj),
            text
        )
